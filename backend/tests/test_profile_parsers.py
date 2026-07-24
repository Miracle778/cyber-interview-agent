from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from app.profile.errors import (
    ProfileEncryptedDocument,
    ProfileNoExtractableText,
    ProfileParseError,
)
from app.profile.parsers import ExtractedSegment, parse_document


def _make_pdf(pages_text: list[str]) -> bytes:
    """Build a minimal valid PDF with one text object per page."""
    num_pages = len(pages_text)
    font_obj_num = 3 + num_pages * 2
    page_obj_nums = [3 + i * 2 for i in range(num_pages)]
    content_obj_nums = [4 + i * 2 for i in range(num_pages)]

    kids = " ".join(f"{n} 0 R" for n in page_obj_nums)
    obj_defs: list[tuple[int, bytes]] = []
    obj_defs.append((1, b"<</Type/Catalog/Pages 2 0 R>>"))
    obj_defs.append(
        (2, b"<</Type/Pages/Kids[" + kids.encode() + b"]/Count " + str(num_pages).encode() + b">>")
    )
    for i in range(num_pages):
        page_obj = (
            b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents "
            + str(content_obj_nums[i]).encode()
            + b" 0 R/Resources<</Font<</F1 "
            + str(font_obj_num).encode()
            + b" 0 R>>>>>>"
        )
        stream = (
            b"BT /F1 12 Tf 72 700 Td ("
            + pages_text[i].encode("latin-1")
            + b") Tj ET"
        )
        content_obj = (
            b"<</Length " + str(len(stream)).encode() + b">>\nstream\n" + stream + b"\nendstream"
        )
        obj_defs.append((page_obj_nums[i], page_obj))
        obj_defs.append((content_obj_nums[i], content_obj))
    obj_defs.append((font_obj_num, b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>"))

    header = b"%PDF-1.4\n"
    body = b""
    offsets: dict[int, int] = {}
    pos = len(header)
    for num, data in obj_defs:
        offsets[num] = pos
        obj_line = f"{num} 0 obj\n".encode() + data + b"\nendobj\n"
        body += obj_line
        pos += len(obj_line)

    xref_pos = len(header) + len(body)
    num_objs = font_obj_num + 1
    xref = f"xref\n0 {num_objs}\n".encode()
    xref += b"0000000000 65535 f \n"
    for num in range(1, num_objs):
        xref += f"{offsets[num]:010d} 00000 n \n".encode()
    trailer = (
        f"trailer\n<</Size {num_objs}/Root 1 0 R>>\nstartxref\n{xref_pos}\n%%EOF".encode()
    )
    return header + body + xref + trailer


def _make_docx(paragraphs: list[str]) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_semantic_docx() -> bytes:
    document = Document()
    document.add_heading("项目经历", level=1)
    document.add_paragraph("Cyber Interview Agent")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_parse_pdf_returns_page_and_page_line_locators() -> None:
    content = _make_pdf(["Page one content", "Page two content"])

    segments = parse_document(
        file_name="resume.pdf", mime_type="application/pdf", content=content
    )

    assert len(segments) == 2
    assert segments[0].text == "Page one content"
    assert segments[0].locator == {
        "page": 1,
        "pageLineStart": 1,
        "pageLineEnd": 1,
    }
    assert segments[1].locator["page"] == 2


def test_parse_docx_returns_paragraph_locators() -> None:
    content = _make_docx(["First paragraph", "Second paragraph"])

    segments = parse_document(
        file_name="resume.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=content,
    )

    assert [s.text for s in segments] == ["First paragraph", "Second paragraph"]
    assert segments[0].locator == {"paragraph": 1}
    assert segments[1].locator == {"paragraph": 2}


def test_parse_markdown_returns_line_locators() -> None:
    content = b"# Title\n\nBody line one\nBody line two\n"

    segments = parse_document(
        file_name="resume.md", mime_type="text/markdown", content=content
    )

    assert len(segments) >= 1
    assert all({"lineStart", "lineEnd"} <= set(s.locator) for s in segments)
    assert "Body line one" in "\n".join(s.text for s in segments)
    assert segments[0].locator["section"] == "Title"
    assert segments[0].locator["block"] == "Title"


def test_parse_markdown_keeps_section_context_for_multiple_blocks() -> None:
    content = (
        "# 项目经历\n\n"
        "## Cyber Interview Agent\n"
        "负责 Agent Runtime 与恢复机制。\n\n"
        "补充了 SSE 状态推送。\n"
    ).encode()

    segments = parse_document(
        file_name="resume.md", mime_type="text/markdown", content=content
    )

    assert [segment.locator.get("section") for segment in segments] == [
        "项目经历",
        "项目经历",
    ]
    assert segments[0].locator["block"] == "Cyber Interview Agent"
    assert segments[1].locator["block"] == "Cyber Interview Agent"


def test_parse_docx_combines_heading_with_first_section_paragraph() -> None:
    segments = parse_document(
        file_name="resume.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=_make_semantic_docx(),
    )

    assert len(segments) == 1
    assert segments[0].text == "项目经历\nCyber Interview Agent"
    assert segments[0].locator == {
        "paragraphStart": 1,
        "paragraphEnd": 2,
        "section": "项目经历",
        "block": "Cyber Interview Agent",
    }


def test_parse_text_normalizes_line_endings() -> None:
    content = b"line1\r\nline2\r\n"

    segments = parse_document(
        file_name="resume.txt", mime_type="text/plain", content=content
    )

    joined = "\n".join(s.text for s in segments if s.text)
    assert "\r" not in joined
    assert "line1" in joined and "line2" in joined


@pytest.mark.parametrize(
    ("file_name", "mime_type", "content"),
    [
        ("empty.txt", "text/plain", b""),
        ("empty.md", "text/markdown", b"\n  \n"),
        (
            "empty.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _make_docx(["", "  "]),
        ),
    ],
)
def test_parse_empty_document_raises_no_extractable_text(
    file_name: str, mime_type: str, content: bytes
) -> None:
    with pytest.raises(ProfileNoExtractableText) as caught:
        parse_document(file_name=file_name, mime_type=mime_type, content=content)

    assert caught.value.code == "profile_no_extractable_text"


def test_parse_corrupt_pdf_raises_redacted_error() -> None:
    with pytest.raises(ProfileParseError) as caught:
        parse_document(
            file_name="bad.pdf", mime_type="application/pdf", content=b"%PDF-1.4 broken"
        )
    assert caught.value.code == "profile_parse_failed"
    assert "broken" not in str(caught.value)


def test_parse_encrypted_pdf_raises_encrypted_error() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt("secret-password")
    buffer = io.BytesIO()
    writer.write(buffer)
    content = buffer.getvalue()

    with pytest.raises(ProfileEncryptedDocument) as caught:
        parse_document(
            file_name="locked.pdf", mime_type="application/pdf", content=content
        )
    assert caught.value.code == "profile_encrypted_document"


def test_parse_scanned_pdf_without_text_raises_recoverable_error() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buffer = io.BytesIO()
    writer.write(buffer)
    content = buffer.getvalue()

    with pytest.raises(ProfileParseError) as caught:
        parse_document(
            file_name="scanned.pdf", mime_type="application/pdf", content=content
        )
    assert caught.value.code == "profile_no_extractable_text"


def test_parse_docx_skips_empty_paragraphs() -> None:
    content = _make_docx(["kept", "", "also kept"])

    segments = parse_document(
        file_name="resume.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=content,
    )

    assert [s.text for s in segments] == ["kept", "also kept"]
