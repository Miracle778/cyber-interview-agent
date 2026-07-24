from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path

from app.profile.errors import (
    ProfileEncryptedDocument,
    ProfileNoExtractableText,
    ProfileParseError,
    ProfileUnsupportedFileType,
)

_PDF_EXT = ".pdf"
_DOCX_EXT = ".docx"
_MARKDOWN_EXTS = {".md", ".markdown"}
_TEXT_EXT = ".txt"
_MARKDOWN_HEADING = re.compile(r"^\s{0,3}#{1,6}\s+(.+?)\s*#*\s*$")
_DECORATED_HEADING = re.compile(r"^\s*(?:\*\*|__)(.+?)(?:\*\*|__)\s*$")
_SECTION_HEADINGS = re.compile(
    r"^(?:"
    r"个人(?:信息|简介|概况|总结)|联系方式|求职意向|"
    r"专业技能|技能(?:清单|特长)?|技术栈|"
    r"项目(?:经历|经验|实践)?|工作(?:经历|经验)|实习(?:经历|经验)|"
    r"教育(?:经历|背景)?|校园经历|"
    r"证书|认证|获奖(?:经历)?|荣誉|成果|个人链接|"
    r"profile|summary|objective|contact|skills?|technical skills?|"
    r"projects?|project experience|work experience|employment|experience|"
    r"education|certifications?|awards?|achievements?|links?"
    r")$",
    re.IGNORECASE,
)
_DETAIL_SECTION_HEADINGS = re.compile(
    r"^(?:项目(?:经历|经验|实践)?|工作(?:经历|经验)|实习(?:经历|经验)|"
    r"教育(?:经历|背景)?|校园经历|projects?|project experience|"
    r"work experience|employment|experience|education)$",
    re.IGNORECASE,
)


@dataclass(frozen=True, slots=True)
class ExtractedSegment:
    text: str
    locator: dict[str, int | str]


def parse_document(
    *, file_name: str, mime_type: str, content: bytes
) -> tuple[ExtractedSegment, ...]:
    """Parse uploaded material bytes into evidence-grounded segments.

    The locator identifies where each segment came from: page/page-line for
    PDF, paragraph/paragraph-range for DOCX, and line range for Markdown/text.
    When recognizable headings exist, ``section`` and ``block`` preserve their
    user-facing meaning. Line endings are normalized in extracted text; source
    bytes are never rewritten. Errors are redacted and use stable codes.
    """
    del mime_type  # dispatch by validated extension; never trust the browser media type
    extension = Path(file_name).suffix.lower()
    if extension == _PDF_EXT:
        return _parse_pdf(content)
    if extension == _DOCX_EXT:
        return _parse_docx(content)
    if extension in _MARKDOWN_EXTS or extension == _TEXT_EXT:
        return _parse_text(content)
    raise ProfileUnsupportedFileType("unsupported file type")


def _parse_pdf(content: bytes) -> tuple[ExtractedSegment, ...]:
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError as error:  # pragma: no cover - dependency guarded by lock
        raise ProfileParseError("pdf parser unavailable") from error

    try:
        reader = PdfReader(io.BytesIO(content))
    except PdfReadError as error:
        raise ProfileParseError("pdf could not be read") from error
    except Exception as error:
        raise ProfileParseError("pdf could not be read") from error

    if reader.is_encrypted:
        try:
            decrypted = reader.decrypt("")
        except Exception:
            decrypted = 0
        if not decrypted:
            raise ProfileEncryptedDocument("document is encrypted")

    segments: list[ExtractedSegment] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        cleaned = text.strip()
        if cleaned:
            segments.extend(
                _segments_from_lines(
                    cleaned.splitlines(),
                    base_locator={"page": index},
                    line_locator_prefix="pageLine",
                )
            )

    if not segments:
        raise ProfileNoExtractableText(
            "pdf has no extractable text (scanned image resumes require OCR, "
            "which is outside the R3 first milestone"
        )
    return tuple(segments)


def _parse_docx(content: bytes) -> tuple[ExtractedSegment, ...]:
    try:
        from docx import Document
    except ImportError as error:  # pragma: no cover - dependency guarded by lock
        raise ProfileParseError("docx parser unavailable") from error

    try:
        document = Document(io.BytesIO(content))
    except Exception as error:
        raise ProfileParseError("docx could not be read") from error

    segments: list[ExtractedSegment] = []
    current_section: str | None = None
    current_block: str | None = None
    pending_heading: tuple[str, int, str] | None = None
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        heading = _heading_text(
            text,
            styled_heading=paragraph.style.name.lower().startswith("heading"),
        )
        if heading:
            if _is_broad_section(heading):
                current_section = heading
                current_block = heading
                pending_heading = (text, index, heading)
            elif pending_heading is not None and current_section:
                pending_text, pending_index, _ = pending_heading
                current_block = heading
                pending_heading = (
                    f"{pending_text}\n{text}",
                    pending_index,
                    heading,
                )
            else:
                current_section = current_section or heading
                current_block = heading
                pending_heading = (text, index, heading)
            continue

        locator: dict[str, int | str]
        segment_text = text
        if pending_heading is not None:
            heading_text, heading_index, heading_block = pending_heading
            segment_text = f"{heading_text}\n{text}"
            section = current_section or _clean_title(heading_text)
            locator = {
                "paragraphStart": heading_index,
                "paragraphEnd": index,
                "section": section,
                "block": (
                    heading_block
                    if heading_block != section
                    else _detail_block_title(section, text) or section
                ),
            }
            pending_heading = None
        else:
            locator = {"paragraph": index}
            _add_semantic_locator(
                locator,
                [text],
                current_section,
                preferred_block=current_block,
            )
        segments.append(ExtractedSegment(text=segment_text, locator=locator))
    if not segments:
        raise ProfileNoExtractableText("document has no extractable text")
    return tuple(segments)


def _parse_text(content: bytes) -> tuple[ExtractedSegment, ...]:
    try:
        raw = content.decode("utf-8")
    except UnicodeDecodeError as error:
        raise ProfileParseError("text must be utf-8 encoded") from error
    normalized = raw.replace("\r\n", "\n").replace("\r", "\n")
    segments = list(
        _segments_from_lines(
            normalized.split("\n"),
            base_locator={},
            line_locator_prefix="line",
        )
    )
    if not segments:
        raise ProfileNoExtractableText("document has no extractable text")
    return tuple(segments)


def _segments_from_lines(
    lines: list[str],
    *,
    base_locator: dict[str, int | str],
    line_locator_prefix: str,
) -> tuple[ExtractedSegment, ...]:
    """Split on document structure while retaining the nearest semantic heading."""
    segments: list[ExtractedSegment] = []
    current_section: str | None = None
    current_block: str | None = None
    start: int | None = None
    buffer: list[str] = []

    def flush(end: int) -> None:
        nonlocal start, buffer
        if start is None or not buffer:
            return
        if len(buffer) == 1 and _heading_text(buffer[0].strip()):
            start = None
            buffer = []
            return
        locator = dict(base_locator)
        locator[f"{line_locator_prefix}Start"] = start
        locator[f"{line_locator_prefix}End"] = end
        _add_semantic_locator(
            locator,
            buffer,
            current_section,
            preferred_block=current_block,
        )
        segments.append(
            ExtractedSegment(text="\n".join(buffer).strip(), locator=locator)
        )
        start = None
        buffer = []

    for index, line in enumerate(lines, start=1):
        stripped = line.strip()
        heading = _heading_text(stripped)
        if heading:
            if (
                buffer
                and all(_heading_text(item.strip()) for item in buffer)
                and current_section
                and not _is_broad_section(heading)
            ):
                current_block = heading
                buffer.append(line)
                continue
            flush(index - 1)
            if _is_broad_section(heading):
                current_section = heading
            elif current_section is None:
                current_section = heading
            current_block = heading
            start = index
            buffer = [line]
            continue
        if stripped:
            if start is None:
                start = index
            buffer.append(line)
            continue
        if buffer:
            # Keep a standalone heading with the first paragraph that follows it.
            if len(buffer) == 1 and _heading_text(buffer[0].strip()):
                continue
            flush(index - 1)
    flush(len(lines))
    return tuple(segments)


def _add_semantic_locator(
    locator: dict[str, int | str],
    lines: list[str],
    current_section: str | None,
    *,
    preferred_block: str | None = None,
) -> None:
    if current_section:
        locator["section"] = current_section
    title = preferred_block or _block_title(lines, current_section)
    if title:
        locator["block"] = title


def _heading_text(value: str, *, styled_heading: bool = False) -> str | None:
    markdown = _MARKDOWN_HEADING.match(value)
    if markdown:
        return _clean_title(markdown.group(1))
    decorated = _DECORATED_HEADING.match(value)
    candidate = _clean_title(decorated.group(1) if decorated else value)
    normalized = candidate.rstrip("：:").strip()
    if styled_heading or _SECTION_HEADINGS.fullmatch(normalized):
        return normalized
    return None


def _block_title(lines: list[str], section: str | None) -> str | None:
    first = next((line.strip() for line in lines if line.strip()), "")
    heading = _heading_text(first)
    if heading:
        return heading
    decorated = _DECORATED_HEADING.match(first)
    if decorated:
        return _clean_title(decorated.group(1))
    if section:
        return _detail_block_title(section, first) or section
    return section


def _detail_block_title(section: str, value: str) -> str | None:
    candidate = _clean_title(value.splitlines()[0])
    if (
        _DETAIL_SECTION_HEADINGS.fullmatch(section)
        and candidate
        and len(candidate) <= 60
        and not candidate.startswith(("-", "*", "•"))
        and not candidate.endswith(("。", "；", "，", ".", ";", ","))
    ):
        return candidate
    return None


def _is_broad_section(value: str) -> bool:
    return _SECTION_HEADINGS.fullmatch(value.rstrip("：:").strip()) is not None


def _clean_title(value: str) -> str:
    cleaned = value.strip()
    cleaned = re.sub(r"^\s{0,3}#{1,6}\s+", "", cleaned)
    cleaned = re.sub(r"^\s*[-*•]\s+", "", cleaned)
    cleaned = re.sub(r"^(?:\*\*|__)(.*?)(?:\*\*|__)$", r"\1", cleaned)
    return cleaned.rstrip("#：:").strip()
